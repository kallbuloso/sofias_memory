from __future__ import annotations

import ipaddress
from io import BytesIO
from pathlib import Path
from typing import cast
from uuid import UUID, uuid4

import httpx
import pytest
from docx import Document as WordDocument
from fastapi import UploadFile
from pypdf import PdfWriter

from sofias_memory.api.errors import SofiasMemoryError
from sofias_memory.api.middleware import API_KEY_HEADER
from sofias_memory.api.routes.remember import parse_metadata_json, read_upload_file_bytes
from sofias_memory.app import create_app
from sofias_memory.config import Settings
from sofias_memory.domain import DatasetStatus, PipelineRunStatus, SourceKind, SourceStatus
from sofias_memory.infrastructure.postgres.models import Dataset, Document, PipelineRun, Source
from sofias_memory.loaders.text import (
    CSV_FILE_MIME_TYPE,
    DOCX_FILE_MIME_TYPE,
    HTML_FILE_MIME_TYPE,
    JSON_FILE_MIME_TYPE,
    MARKDOWN_FILE_MIME_TYPE,
    PDF_FILE_MIME_TYPE,
    TEXT_FILE_MIME_TYPE,
    TextFileLoadError,
    prepare_text_file_content,
)
from sofias_memory.loaders.url import FetchedUrlContent, fetch_https_url
from sofias_memory.schemas.remember import (
    RememberTextRequest,
    RememberTextResult,
    RememberUrlRequest,
)
from sofias_memory.services.remember import RememberService, RememberUnitOfWork, UnitOfWorkFactory

EXPECTED_API_KEY = "sf-AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA"
DATABASE_URL = "postgresql+asyncpg://sofias_memory:fake@postgres:5432/sofias_memory"
NEO4J_PASSWORD = "fake-neo4j-password"
LLM_API_KEY = "sk-fake-test-key"


def make_settings(tmp_path: Path) -> Settings:
    return Settings(
        _env_file=None,  # type: ignore[call-arg]
        api_key=EXPECTED_API_KEY,
        database_url=DATABASE_URL,
        neo4j_password=NEO4J_PASSWORD,
        llm_api_key=LLM_API_KEY,
        app_env="test",
        data_directory=tmp_path,
    )


class FakeStore:
    def __init__(self) -> None:
        self.datasets: list[Dataset] = []
        self.sources: list[Source] = []
        self.documents: list[Document] = []
        self.pipeline_runs: list[PipelineRun] = []
        self.commits = 0


class FakeDatasetRepository:
    def __init__(self, store: FakeStore) -> None:
        self._store = store

    async def add(self, dataset: Dataset) -> Dataset:
        self._store.datasets.append(dataset)
        return dataset

    async def get_by_slug(self, slug: str) -> Dataset | None:
        return next((dataset for dataset in self._store.datasets if dataset.slug == slug), None)


class FakeSourceRepository:
    def __init__(self, store: FakeStore) -> None:
        self._store = store

    async def add(self, source: Source) -> Source:
        self._store.sources.append(source)
        return source

    async def get_latest_by_content_hash(
        self,
        *,
        dataset_id: UUID,
        content_sha256: str,
    ) -> Source | None:
        matches = [
            source
            for source in self._store.sources
            if source.dataset_id == dataset_id and source.content_sha256 == content_sha256
        ]
        return max(matches, key=lambda source: source.version, default=None)


class FakeDocumentRepository:
    def __init__(self, store: FakeStore) -> None:
        self._store = store

    async def add(self, document: Document) -> Document:
        self._store.documents.append(document)
        return document

    async def list_for_source(self, source_id: UUID) -> list[Document]:
        return [document for document in self._store.documents if document.source_id == source_id]


class FakePipelineRunRepository:
    def __init__(self, store: FakeStore) -> None:
        self._store = store

    async def add(self, run: PipelineRun) -> PipelineRun:
        self._store.pipeline_runs.append(run)
        return run

    async def get_by_id(self, run_id: UUID) -> PipelineRun | None:
        return next((run for run in self._store.pipeline_runs if run.id == run_id), None)

    async def get_by_idempotency_key(self, idempotency_key: str) -> PipelineRun | None:
        return next(
            (run for run in self._store.pipeline_runs if run.idempotency_key == idempotency_key),
            None,
        )


class FakeUnitOfWork:
    def __init__(self, store: FakeStore) -> None:
        self._store = store
        self.datasets = FakeDatasetRepository(store)
        self.sources = FakeSourceRepository(store)
        self.documents = FakeDocumentRepository(store)
        self.pipeline_runs = FakePipelineRunRepository(store)

    async def __aenter__(self) -> FakeUnitOfWork:
        return self

    async def __aexit__(self, exc_type: object, exc: object, traceback: object) -> None:
        return None

    async def commit(self) -> None:
        self._store.commits += 1


def service_for(tmp_path: Path, store: FakeStore) -> RememberService:
    def create_uow() -> RememberUnitOfWork:
        return cast(RememberUnitOfWork, FakeUnitOfWork(store))

    return RememberService(
        make_settings(tmp_path),
        unit_of_work_factory=cast(UnitOfWorkFactory, create_uow),
    )


def remember_request(
    content: str = "Sofias Memory mantem memoria persistente.",
    *,
    force: bool = False,
    name: str = "nota-inicial",
) -> RememberTextRequest:
    return RememberTextRequest(
        dataset="main",
        content=content,
        name=name,
        metadata={"origin": "unit"},
        session_id="chat-42",
        mode="ingest",
        wait=True,
        force=force,
    )


def make_text_pdf(*page_texts: str) -> bytes:
    objects: list[bytes] = []
    page_object_numbers: list[int] = []
    content_object_numbers: list[int] = []
    font_object_number = 3 + (len(page_texts) * 2)

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"")
    for page_text in page_texts:
        page_object_number = len(objects) + 1
        content_object_number = page_object_number + 1
        page_object_numbers.append(page_object_number)
        content_object_numbers.append(content_object_number)
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_object_number} 0 R >> >> "
                f"/Contents {content_object_number} 0 R >>"
            ).encode("ascii")
        )
        stream = f"BT /F1 12 Tf 72 720 Td ({pdf_literal(page_text)}) Tj ET".encode("ascii")
        objects.append(
            b"<< /Length "
            + str(len(stream)).encode("ascii")
            + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )

    kids = " ".join(f"{page_number} 0 R" for page_number in page_object_numbers)
    objects[1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_object_numbers)} >>".encode(
        "ascii"
    )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    return render_pdf_objects(objects)


def pdf_literal(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def render_pdf_objects(objects: list[bytes]) -> bytes:
    chunks = [b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"]
    offsets = [0]
    offset = len(chunks[0])
    for index, body in enumerate(objects, start=1):
        obj = f"{index} 0 obj\n".encode("ascii") + body + b"\nendobj\n"
        offsets.append(offset)
        chunks.append(obj)
        offset += len(obj)
    xref_offset = offset
    xref = [f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("ascii")]
    xref.extend(f"{object_offset:010d} 00000 n \n".encode("ascii") for object_offset in offsets[1:])
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    ).encode("ascii")
    return b"".join([*chunks, *xref, trailer])


def make_blank_pdf() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(output)
    return output.getvalue()


def make_docx_with_content() -> bytes:
    document = WordDocument()
    document.add_heading("Heading", level=1)
    document.add_paragraph("Intro paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Sofias"
    table.cell(1, 1).text = "Memory"
    document.add_paragraph("After table.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_empty_docx() -> bytes:
    output = BytesIO()
    WordDocument().save(output)
    return output.getvalue()


@pytest.mark.asyncio
async def test_remember_text_creates_source_document_and_default_dataset(tmp_path: Path) -> None:
    store = FakeStore()
    result = await service_for(tmp_path, store).remember_text(remember_request())

    assert result.status == PipelineRunStatus.SUCCEEDED.value
    assert result.deduplicated is False
    assert len(store.datasets) == 1
    assert store.datasets[0].slug == "main"
    assert store.datasets[0].status == DatasetStatus.ACTIVE
    assert len(store.sources) == 1
    assert store.sources[0].status == SourceStatus.PENDING
    assert store.sources[0].version == 1
    assert len(store.documents) == 1
    assert store.documents[0].language == "und"
    assert store.documents[0].metadata_["session_id"] == "chat-42"
    assert store.documents[0].generation == store.datasets[0].active_generation
    assert store.pipeline_runs[0].metrics["remember_result"]["source_id"] == str(result.source_id)

    stored_file = tmp_path / str(result.dataset_id) / str(result.source_id) / "original.txt"
    assert stored_file.read_text(encoding="utf-8") == remember_request().content


@pytest.mark.asyncio
async def test_remember_text_deduplicates_same_content_without_force(tmp_path: Path) -> None:
    store = FakeStore()
    service = service_for(tmp_path, store)

    first = await service.remember_text(remember_request())
    second = await service.remember_text(remember_request(name="another-name"))

    assert second.deduplicated is True
    assert second.source_id == first.source_id
    assert second.document_id == first.document_id
    assert len(store.sources) == 1
    assert len(store.documents) == 1
    assert len(store.pipeline_runs) == 2


@pytest.mark.asyncio
async def test_remember_text_force_creates_new_source_version(tmp_path: Path) -> None:
    store = FakeStore()
    service = service_for(tmp_path, store)

    first = await service.remember_text(remember_request())
    second = await service.remember_text(remember_request(force=True))

    assert second.deduplicated is False
    assert second.source_id != first.source_id
    assert [source.version for source in store.sources] == [1, 2]
    assert len(store.documents) == 2


@pytest.mark.asyncio
async def test_idempotency_key_rejects_same_key_with_different_payload(tmp_path: Path) -> None:
    store = FakeStore()
    service = service_for(tmp_path, store)

    first = await service.remember_text(remember_request("first"), idempotency_key="same-key")
    replay = await service.remember_text(remember_request("first"), idempotency_key="same-key")

    assert replay == first
    assert len(store.pipeline_runs) == 1

    with pytest.raises(SofiasMemoryError) as exc_info:
        await service.remember_text(remember_request("second"), idempotency_key="same-key")

    assert exc_info.value.status_code == 409
    assert len(store.sources) == 1
    assert len(store.documents) == 1


@pytest.mark.asyncio
async def test_remember_txt_file_creates_source_document_and_storage(tmp_path: Path) -> None:
    store = FakeStore()
    original_bytes = b"Sofias Memory\r\nfile ingest."
    prepared_file = prepare_text_file_content("note.txt", original_bytes)

    result = await service_for(tmp_path, store).remember_file(
        dataset="main",
        prepared_file=prepared_file,
        metadata={"origin": "upload"},
        session_id="file-session",
        mode="ingest",
        wait=True,
        force=False,
    )

    source = store.sources[0]
    document = store.documents[0]
    assert source.kind == SourceKind.FILE
    assert source.mime_type == TEXT_FILE_MIME_TYPE
    assert source.name == "note.txt"
    assert source.content_sha256 == result.content_hash
    assert source.normalized_sha256 == document.text_sha256
    assert source.content_sha256 != source.normalized_sha256
    assert document.normalized_text == "Sofias Memory\nfile ingest."
    assert document.metadata_["session_id"] == "file-session"

    stored_file = tmp_path / str(result.dataset_id) / str(result.source_id) / "original.txt"
    assert stored_file.read_bytes() == original_bytes


@pytest.mark.asyncio
async def test_remember_markdown_preserves_headings_and_removes_nul(tmp_path: Path) -> None:
    store = FakeStore()
    original_bytes = b"\xef\xbb\xbf# Title\r\n\nBody\x00 text"
    prepared_file = prepare_text_file_content("../note.markdown", original_bytes)

    await service_for(tmp_path, store).remember_file(
        dataset="main",
        prepared_file=prepared_file,
        metadata={},
        session_id=None,
        mode="ingest",
        wait=True,
        force=False,
    )

    assert store.sources[0].kind == SourceKind.FILE
    assert store.sources[0].mime_type == MARKDOWN_FILE_MIME_TYPE
    assert store.sources[0].name == "note.markdown"
    assert store.documents[0].title == "note.markdown"
    assert store.documents[0].normalized_text == "# Title\n\nBody text"


def test_json_file_is_validated_and_normalized_deterministically() -> None:
    original_bytes = b'{"z":2,"a":{"b":1},"list":[true,null]}'
    prepared_file = prepare_text_file_content("payload.json", original_bytes)

    assert prepared_file.mime_type == JSON_FILE_MIME_TYPE
    assert prepared_file.storage_extension == ".json"
    assert prepared_file.text.original_bytes == original_bytes
    assert prepared_file.text.normalized_text == (
        '{\n  "a": {\n    "b": 1\n  },\n  "list": [\n    true,\n    null\n  ],\n  "z": 2\n}\n'
    )

    with pytest.raises(TextFileLoadError, match="valid JSON"):
        prepare_text_file_content("broken.json", b'{"a":')


def test_csv_file_is_validated_and_preserved() -> None:
    original_bytes = b"name,value\r\nSofias,1\r\nMemory,2"
    prepared_file = prepare_text_file_content("table.csv", original_bytes)

    assert prepared_file.mime_type == CSV_FILE_MIME_TYPE
    assert prepared_file.storage_extension == ".csv"
    assert prepared_file.text.original_bytes == original_bytes
    assert prepared_file.text.normalized_text == "name,value\nSofias,1\nMemory,2"

    with pytest.raises(TextFileLoadError, match="must not be empty"):
        prepare_text_file_content("empty.csv", b"")
    with pytest.raises(TextFileLoadError, match="valid CSV"):
        prepare_text_file_content("broken.csv", b'"unterminated')


def test_html_file_extracts_visible_text_and_ignores_non_visible_content() -> None:
    original_bytes = (
        b"<html><head><style>.hidden{}</style><script>alert(1)</script></head>"
        b"<body><h1>Title</h1><p>Hello <strong>world</strong>.</p>"
        b"<noscript>ignore me</noscript><ul><li>First</li><li>Second</li></ul></body></html>"
    )
    prepared_file = prepare_text_file_content("page.htm", original_bytes)

    assert prepared_file.mime_type == HTML_FILE_MIME_TYPE
    assert prepared_file.storage_extension == ".html"
    assert prepared_file.text.normalized_text == "Title\nHello world.\nFirst\nSecond"

    html_file = prepare_text_file_content("page.html", b"<div>Visible<br>Text</div>")
    assert html_file.mime_type == HTML_FILE_MIME_TYPE
    assert html_file.text.normalized_text == "Visible\nText"

    with pytest.raises(TextFileLoadError, match="visible text"):
        prepare_text_file_content("empty.html", b"<script>onlyHidden()</script><style>x{}</style>")


def test_pdf_file_extracts_textual_pages_and_preserves_original_bytes() -> None:
    original_bytes = make_text_pdf("First page", "Second page")
    prepared_file = prepare_text_file_content("document.pdf", original_bytes)

    assert prepared_file.mime_type == PDF_FILE_MIME_TYPE
    assert prepared_file.storage_extension == ".pdf"
    assert prepared_file.text.original_bytes == original_bytes
    assert prepared_file.text.normalized_text == "First page\n\nSecond page"


def test_pdf_file_rejects_invalid_or_non_textual_pdf() -> None:
    with pytest.raises(TextFileLoadError, match="readable textual PDF"):
        prepare_text_file_content("broken.pdf", b"not a pdf")

    with pytest.raises(TextFileLoadError, match="extractable text"):
        prepare_text_file_content("blank.pdf", make_blank_pdf())


def test_docx_file_extracts_body_text_tables_and_preserves_order() -> None:
    original_bytes = make_docx_with_content()
    prepared_file = prepare_text_file_content("document.docx", original_bytes)

    assert prepared_file.mime_type == DOCX_FILE_MIME_TYPE
    assert prepared_file.storage_extension == ".docx"
    assert prepared_file.text.original_bytes == original_bytes
    assert prepared_file.text.normalized_text == (
        "Heading\nIntro paragraph.\nName\tValue\nSofias\tMemory\nAfter table."
    )


def test_docx_file_rejects_invalid_or_empty_docx() -> None:
    with pytest.raises(TextFileLoadError, match="readable DOCX"):
        prepare_text_file_content("broken.docx", b"not a docx")

    with pytest.raises(TextFileLoadError, match="must contain text"):
        prepare_text_file_content("empty.docx", make_empty_docx())


@pytest.mark.asyncio
async def test_remember_file_deduplicates_and_force_creates_next_version(tmp_path: Path) -> None:
    store = FakeStore()
    service = service_for(tmp_path, store)
    prepared_file = prepare_text_file_content("note.md", b"# Same\n")

    first = await service.remember_file(
        dataset="main",
        prepared_file=prepared_file,
        metadata={},
        session_id=None,
        mode="ingest",
        wait=True,
        force=False,
    )
    duplicate = await service.remember_file(
        dataset="main",
        prepared_file=prepared_file,
        metadata={},
        session_id=None,
        mode="ingest",
        wait=True,
        force=False,
    )
    forced = await service.remember_file(
        dataset="main",
        prepared_file=prepared_file,
        metadata={},
        session_id=None,
        mode="ingest",
        wait=True,
        force=True,
    )

    assert duplicate.deduplicated is True
    assert duplicate.source_id == first.source_id
    assert forced.deduplicated is False
    assert forced.source_id != first.source_id
    assert [source.version for source in store.sources] == [1, 2]


@pytest.mark.asyncio
async def test_remember_file_rejects_unsupported_metadata_and_oversize() -> None:
    with pytest.raises(TextFileLoadError, match="Unsupported"):
        prepare_text_file_content("data.docm", b"not supported")

    with pytest.raises(SofiasMemoryError) as metadata_error:
        parse_metadata_json("[1, 2, 3]")
    assert metadata_error.value.status_code == 400

    upload = UploadFile(filename="large.txt", file=cast(object, _BytesFile(b"abcdef")))
    with pytest.raises(SofiasMemoryError) as size_error:
        await read_upload_file_bytes(upload, max_bytes=3)
    assert size_error.value.status_code == 413


@pytest.mark.asyncio
async def test_remember_file_route_accepts_new_formats(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    seen_mime_types: list[str] = []

    class FakeRememberService:
        def __init__(self, *args: object, **kwargs: object) -> None:
            pass

        async def remember_file(self, **kwargs: object) -> RememberTextResult:
            prepared_file = kwargs["prepared_file"]
            seen_mime_types.append(prepared_file.mime_type)
            return RememberTextResult(
                run_id=uuid4(),
                status=PipelineRunStatus.SUCCEEDED.value,
                dataset_id=uuid4(),
                source_id=uuid4(),
                document_id=uuid4(),
                content_hash=prepared_file.text.content_sha256,
                chunks=0,
                entities=0,
                relations=0,
                deduplicated=False,
            )

    monkeypatch.setattr("sofias_memory.api.routes.remember.RememberService", FakeRememberService)
    app = create_app(make_settings(tmp_path), enable_postgres_readiness=False, enable_neo4j=False)
    transport = httpx.ASGITransport(app=app)

    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        for filename, content in (
            ("payload.json", b'{"ok": true}'),
            ("table.csv", b"a,b\n1,2\n"),
            ("page.html", b"<p>Hello</p>"),
            ("document.pdf", make_text_pdf("PDF text")),
            ("document.docx", make_docx_with_content()),
        ):
            response = await client.post(
                "/api/v1/remember/file",
                headers={API_KEY_HEADER: EXPECTED_API_KEY},
                files={"file": (filename, content, "application/octet-stream")},
                data={"mode": "ingest", "wait": "true"},
            )
            assert response.status_code == 200

    assert seen_mime_types == [
        JSON_FILE_MIME_TYPE,
        CSV_FILE_MIME_TYPE,
        HTML_FILE_MIME_TYPE,
        PDF_FILE_MIME_TYPE,
        DOCX_FILE_MIME_TYPE,
    ]


@pytest.mark.asyncio
async def test_fetch_https_url_accepts_supported_content_and_redirects() -> None:
    requests: list[str] = []

    async def handler(request: httpx.Request) -> httpx.Response:
        requests.append(str(request.url))
        if request.url.path == "/start":
            return httpx.Response(302, headers={"Location": "/page"})
        if request.url.path == "/document":
            return httpx.Response(
                200,
                headers={"Content-Type": "application/pdf"},
                content=b"%PDF-placeholder",
            )
        return httpx.Response(
            200,
            headers={"Content-Type": "text/html; charset=utf-8"},
            content=b"<h1>Hello</h1>",
        )

    fetched = await fetch_https_url(
        "https://example.com/start#fragment",
        max_bytes=1024,
        transport=httpx.MockTransport(handler),
        resolver=public_resolver,
    )
    pdf_fetched = await fetch_https_url(
        "https://example.com/document",
        max_bytes=1024,
        transport=httpx.MockTransport(handler),
        resolver=public_resolver,
    )

    assert requests[:2] == ["https://example.com/start", "https://example.com/page"]
    assert fetched.requested_url == "https://example.com/start"
    assert fetched.filename == "page.html"
    assert fetched.media_type == HTML_FILE_MIME_TYPE
    assert fetched.body == b"<h1>Hello</h1>"
    assert pdf_fetched.filename == "document.pdf"
    assert pdf_fetched.media_type == PDF_FILE_MIME_TYPE


@pytest.mark.asyncio
async def test_fetch_https_url_blocks_ssrf_and_unsafe_redirects() -> None:
    async def redirect_to_private(request: httpx.Request) -> httpx.Response:
        del request
        return httpx.Response(302, headers={"Location": "https://private.example/page"})

    async def redirect_to_http(request: httpx.Request) -> httpx.Response:
        del request
        return httpx.Response(302, headers={"Location": "http://example.com/page"})

    async def redirect_loop(request: httpx.Request) -> httpx.Response:
        del request
        return httpx.Response(302, headers={"Location": "/loop"})

    with pytest.raises(SofiasMemoryError, match="HTTPS"):
        await fetch_https_url(
            "http://example.com/page",
            max_bytes=1024,
            transport=httpx.MockTransport(redirect_to_http),
            resolver=public_resolver,
        )
    with pytest.raises(SofiasMemoryError, match="credentials"):
        await fetch_https_url(
            "https://user:pass@example.com/page",
            max_bytes=1024,
            transport=httpx.MockTransport(redirect_to_http),
            resolver=public_resolver,
        )
    with pytest.raises(SofiasMemoryError, match="blocked network"):
        await fetch_https_url(
            "https://10.0.0.1/page",
            max_bytes=1024,
            transport=httpx.MockTransport(redirect_to_http),
            resolver=private_resolver,
        )
    with pytest.raises(SofiasMemoryError, match="blocked network"):
        await fetch_https_url(
            "https://example.com/page",
            max_bytes=1024,
            transport=httpx.MockTransport(redirect_to_private),
            resolver=resolver_blocking_private_hostname,
        )
    with pytest.raises(SofiasMemoryError, match="HTTPS"):
        await fetch_https_url(
            "https://example.com/page",
            max_bytes=1024,
            transport=httpx.MockTransport(redirect_to_http),
            resolver=public_resolver,
        )
    with pytest.raises(SofiasMemoryError, match="redirect loop"):
        await fetch_https_url(
            "https://example.com/loop",
            max_bytes=1024,
            transport=httpx.MockTransport(redirect_loop),
            resolver=public_resolver,
        )


@pytest.mark.asyncio
async def test_fetch_https_url_enforces_size_and_supported_content_type() -> None:
    async def content_length_too_large(request: httpx.Request) -> httpx.Response:
        del request
        return httpx.Response(
            200,
            headers={"Content-Type": "text/html", "Content-Length": "5"},
            content=b"short",
        )

    async def streamed_too_large(request: httpx.Request) -> httpx.Response:
        del request
        return httpx.Response(200, headers={"Content-Type": "text/html"}, content=b"abcdef")

    async def unsupported_content_type(request: httpx.Request) -> httpx.Response:
        del request
        return httpx.Response(200, headers={"Content-Type": "application/octet-stream"})

    with pytest.raises(SofiasMemoryError) as content_length_error:
        await fetch_https_url(
            "https://example.com/page",
            max_bytes=4,
            transport=httpx.MockTransport(content_length_too_large),
            resolver=public_resolver,
        )
    assert content_length_error.value.status_code == 413

    with pytest.raises(SofiasMemoryError) as streaming_error:
        await fetch_https_url(
            "https://example.com/page",
            max_bytes=4,
            transport=httpx.MockTransport(streamed_too_large),
            resolver=public_resolver,
        )
    assert streaming_error.value.status_code == 413

    with pytest.raises(SofiasMemoryError, match="content type"):
        await fetch_https_url(
            "https://example.com/page",
            max_bytes=1024,
            transport=httpx.MockTransport(unsupported_content_type),
            resolver=public_resolver,
        )


@pytest.mark.asyncio
async def test_remember_url_creates_url_source_and_idempotent_run(tmp_path: Path) -> None:
    store = FakeStore()
    service = service_for(tmp_path, store)
    request = RememberUrlRequest(
        dataset="main",
        url="https://example.com/page",
        metadata={"origin": "unit"},
        session_id="url-session",
    )
    prepared_file = prepare_text_file_content("page.html", b"<h1>Hello URL</h1>")

    first = await service.remember_url(
        request,
        prepared_file=prepared_file,
        original_uri="https://example.com/page",
        idempotency_key="url-key",
    )
    replay = await service.remember_url(
        request,
        prepared_file=prepared_file,
        original_uri="https://example.com/page",
        idempotency_key="url-key",
    )

    assert replay == first
    assert len(store.sources) == 1
    assert store.sources[0].kind == SourceKind.URL
    assert store.sources[0].original_uri == "https://example.com/page"
    assert store.sources[0].mime_type == HTML_FILE_MIME_TYPE
    assert store.documents[0].normalized_text == "Hello URL"


@pytest.mark.asyncio
async def test_remember_url_idempotency_key_rejects_different_url(tmp_path: Path) -> None:
    store = FakeStore()
    service = service_for(tmp_path, store)
    prepared_file = prepare_text_file_content("page.html", b"<p>Same body</p>")
    first_request = RememberUrlRequest(dataset="main", url="https://example.com/one")
    second_request = RememberUrlRequest(dataset="main", url="https://example.com/two")

    await service.remember_url(
        first_request,
        prepared_file=prepared_file,
        original_uri="https://example.com/one",
        idempotency_key="url-key",
    )

    with pytest.raises(SofiasMemoryError) as exc_info:
        await service.remember_url(
            second_request,
            prepared_file=prepared_file,
            original_uri="https://example.com/two",
            idempotency_key="url-key",
        )

    assert exc_info.value.status_code == 409


@pytest.mark.asyncio
async def test_remember_url_route_uses_fetcher_loader_and_service(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    seen_original_uri: list[str] = []
    seen_mime_types: list[str] = []

    async def fake_fetch(url: str, *, max_bytes: int) -> FetchedUrlContent:
        assert url == "https://example.com/page#ignored"
        assert max_bytes > 0
        return FetchedUrlContent(
            requested_url="https://example.com/page",
            final_url="https://example.com/page",
            filename="page.html",
            media_type=HTML_FILE_MIME_TYPE,
            body=b"<h1>Remote</h1>",
        )

    class FakeRememberService:
        def __init__(self, *args: object, **kwargs: object) -> None:
            pass

        async def remember_url(self, *args: object, **kwargs: object) -> RememberTextResult:
            prepared_file = kwargs["prepared_file"]
            seen_mime_types.append(prepared_file.mime_type)
            seen_original_uri.append(cast(str, kwargs["original_uri"]))
            return RememberTextResult(
                run_id=uuid4(),
                status=PipelineRunStatus.SUCCEEDED.value,
                dataset_id=uuid4(),
                source_id=uuid4(),
                document_id=uuid4(),
                content_hash=prepared_file.text.content_sha256,
                chunks=0,
                entities=0,
                relations=0,
                deduplicated=False,
            )

    monkeypatch.setattr("sofias_memory.api.routes.remember.fetch_https_url", fake_fetch)
    monkeypatch.setattr("sofias_memory.api.routes.remember.RememberService", FakeRememberService)
    app = create_app(make_settings(tmp_path), enable_postgres_readiness=False, enable_neo4j=False)
    transport = httpx.ASGITransport(app=app)

    async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
        response = await client.post(
            "/api/v1/remember/url",
            headers={API_KEY_HEADER: EXPECTED_API_KEY, "Idempotency-Key": "url-route-key"},
            json={"url": "https://example.com/page#ignored", "metadata": {"origin": "unit"}},
        )

    assert response.status_code == 200
    assert seen_original_uri == ["https://example.com/page"]
    assert seen_mime_types == [HTML_FILE_MIME_TYPE]


def public_resolver(host: str, port: int) -> list[ipaddress.IPv4Address | ipaddress.IPv6Address]:
    del host, port
    return [ipaddress.ip_address("93.184.216.34")]


def private_resolver(host: str, port: int) -> list[ipaddress.IPv4Address | ipaddress.IPv6Address]:
    del host, port
    return [ipaddress.ip_address("10.0.0.1")]


def resolver_blocking_private_hostname(
    host: str,
    port: int,
) -> list[ipaddress.IPv4Address | ipaddress.IPv6Address]:
    del port
    if host == "private.example":
        return [ipaddress.ip_address("10.0.0.1")]
    return [ipaddress.ip_address("93.184.216.34")]


class _BytesFile:
    def __init__(self, content: bytes) -> None:
        self._content = content
        self._offset = 0

    def read(self, size: int = -1) -> bytes:
        if size < 0:
            size = len(self._content) - self._offset
        start = self._offset
        self._offset = min(len(self._content), self._offset + size)
        return self._content[start : self._offset]
